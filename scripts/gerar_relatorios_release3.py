#!/usr/bin/env python3
"""Gera relatorio-processo-release3 e relatorio-release3 (.docx)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RELEASES = ROOT / "releases"
DIAGRAMAS = RELEASES / "diagramas"
SEQ = DIAGRAMAS / "documentacao-rf31-rf43"


def set_run_font(run, size=12, bold=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_centered(doc, text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_custom(doc, text, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=12)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    if level:
        p.paragraph_format.left_indent = Cm(1.25 * level)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_placeholder(doc, label: str):
    """Espaço reservado para print de terminal / captura de tela."""
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = sep.add_run("─" * 52)
    set_run_font(r0, size=10)
    r0.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "[ESPAÇO RESERVADO PARA PRINT]\n"
        f"{label}\n"
        "(Inserir captura de tela aqui)"
    )
    set_run_font(run, size=11, bold=True)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sep2.add_run("─" * 52)
    set_run_font(r1, size=10)
    r1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def add_image_if_exists(doc, path: Path, width_cm=16.0, caption: str | None = None):
    if not path.exists():
        add_body(doc, f"[Imagem não encontrada: {path.name}]", first_line_indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=10, bold=True)


def capa(doc, titulo: str):
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "UNIVERSIDADE ESTADUAL DA PARAÍBA", size=14, bold=True)
    add_centered(doc, "CENTRO DE CIÊNCIAS E TECNOLOGIA - CCT", size=12, bold=True)
    add_centered(doc, "DEPARTAMENTO DE COMPUTAÇÃO - DC", size=12, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    for nome in [
        "ALYSSANDRO DYOGO PEREIRA RAMOS",
        "DJHONATAH WESLEY CAVALCANTI ALVES",
        "JOSE LUCAS CUSTODIO SILVA",
        "RODRIGO ALMEIDA GOMES",
    ]:
        add_centered(doc, nome, size=12, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, titulo, size=16, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    add_centered(doc, "CAMPINA GRANDE - PB", size=12, bold=True)
    add_centered(doc, "2026", size=12, bold=True)
    doc.add_page_break()


def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    return doc


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    doc.add_paragraph()


SEQ_ITEMS = [
    ("RF31 — Lançar notas das etapas", ["sequencia-rf31-1.png", "sequencia-rf31-2.png"]),
    ("RF32 — Cálculo automático da média final", ["sequencia-rf32-1.png", "sequencia-rf32-2.png"]),
    ("RF33 — Aluno consulta suas notas", ["sequencia-rf33-1.png", "sequencia-rf33-2.png"]),
    (
        "RF34 — Situação acadêmica do aluno",
        ["sequencia-rf34-1.png", "sequencia-rf34-2.png", "sequencia-rf34-3.png"],
    ),
    (
        "RF35 — Alterar notas antes do fechamento",
        ["sequencia-rf35-1.png", "sequencia-rf35-2.png", "sequencia-rf35-3.png"],
    ),
    (
        "RF36 — Manter histórico das disciplinas cursadas",
        ["sequencia-rf36-1.png", "sequencia-rf36-2.png"],
    ),
    ("RF37 — Campos do histórico acadêmico", ["sequencia-rf37-1.png", "sequencia-rf37-2.png"]),
    (
        "RF38 — Aluno consulta histórico acadêmico",
        ["sequencia-rf38-1.png", "sequencia-rf38-2.png"],
    ),
    (
        "RF39 — Coordenador consulta histórico do curso",
        ["sequencia-rf39-1.png", "sequencia-rf39-2.png", "sequencia-rf39-3.png"],
    ),
    (
        "RF40 — Relatório de alunos matriculados por turma (PDF)",
        ["sequencia-rf40-1.png", "sequencia-rf40-2.png"],
    ),
    (
        "RF41 — Relatório de ocupação de vagas (PDF)",
        ["sequencia-rf41-1.png", "sequencia-rf41-2.png"],
    ),
    (
        "RF42 — Relatório de reprovação por disciplina (PDF)",
        ["sequencia-rf42-1.png", "sequencia-rf42-2.png"],
    ),
    (
        "RF43 — Relatório geral de usuários cadastrados (PDF)",
        ["sequencia-rf43-1.png", "sequencia-rf43-2.png"],
    ),
]


FUNCIONALIDADES = [
    (
        "Lançamento e alteração de notas",
        "O professor lança notas das etapas 1 e 2, altera valores antes do fechamento da turma "
        "e o sistema calcula automaticamente a média final e a situação acadêmica "
        "(aprovado, em recuperação, reprovado por nota ou por falta).",
        "Print — Professor lançando/alterando notas (opções 25/26)",
    ),
    (
        "Consulta de notas pelo aluno",
        "O aluno consulta suas notas por turma, visualizando etapa1, etapa2, média final, "
        "frequência e situação acadêmica.",
        "Print — Aluno consultando notas (opção 25)",
    ),
    (
        "Fechamento de turma e histórico acadêmico",
        "O coordenador fecha a turma, gerando o histórico com período, disciplina, professor, "
        "nota final, frequência e situação. Aluno e coordenador consultam o histórico.",
        "Print — Fechamento de turma / consulta de histórico (opções 25/27/28)",
    ),
    (
        "Relatórios acadêmicos e administrativos em PDF",
        "O coordenador gera relatórios de alunos por turma, ocupação de vagas e reprovação "
        "por disciplina; o administrador gera o relatório geral de usuários. Todos podem "
        "ser baixados em arquivo PDF.",
        "Print — Geração e download de relatórios PDF (opções 29–32)",
    ),
]


def gerar_processo():
    doc = new_doc()
    capa(doc, "RELATÓRIO DE PROCESSO DE RELEASE - 3")

    add_heading_custom(doc, "1. Introdução")
    add_body(
        doc,
        "O seguinte relatório descreve o processo de desenvolvimento, gestão e linha "
        "do tempo de progresso de desenvolvimento das atividades da release 3 da "
        "aplicação ClassRoomPB dividida em 2 sprints às quais ao todo compuseram o "
        "desenvolvimento de 13 requisitos funcionais listados abaixo e divididos em 3 macro "
        "grupos, além de testes das mesmas e documentação.",
    )

    add_bullet(doc, "Notas e Avaliação")
    for item in [
        "RF31: O professor deve poder lançar notas das etapas (etapa1 e etapa2).",
        "RF32: O sistema deve calcular automaticamente a média final.",
        "RF33: O aluno deve poder consultar suas notas.",
        "RF34: O sistema deve informar a situação do aluno: aprovado, reprovado por nota, "
        "reprovado por falta ou em recuperação.",
        "RF35: O professor deve poder alterar notas antes do fechamento da turma.",
    ]:
        add_bullet(doc, item, level=1)

    add_bullet(doc, "Histórico Acadêmico")
    for item in [
        "RF36: O sistema deve manter histórico das disciplinas cursadas pelo aluno.",
        "RF37: O histórico deve registrar período, disciplina, professor, nota final, frequência e situação.",
        "RF38: O aluno deve poder consultar seu histórico acadêmico.",
        "RF39: O coordenador deve poder consultar o histórico dos alunos do curso.",
    ]:
        add_bullet(doc, item, level=1)

    add_bullet(doc, "Relatórios")
    for item in [
        "RF40: O coordenador deve gerar relatório de alunos matriculados por turma.",
        "RF41: O coordenador deve gerar relatório de ocupação de vagas.",
        "RF42: O coordenador deve gerar relatório de reprovação por disciplina.",
        "RF43: O administrador deve gerar relatório geral de usuários cadastrados.",
    ]:
        add_bullet(doc, item, level=1)

    add_heading_custom(doc, "2. Sprint 5")
    add_heading_custom(doc, "2.1 Papéis Assumidos na Sprint 5")
    for item in [
        "Alyssandro Ramos: Desenvolvedor e Documentação",
        "Rodrigo Almeida: Desenvolvedor e Documentação",
        "Djhonatah Wesley: QA",
        "José Lucas: Gerente de Projeto e Desenvolvedor",
    ]:
        add_bullet(doc, item)

    add_heading_custom(doc, "2.2 Backlog da Sprint")
    add_placeholder(
        doc,
        "Backlog da Sprint 5 (RF31–RF36 / quadro do gerenciador de projeto)",
    )

    add_heading_custom(doc, "2.3 Acompanhamento e Desempenho (Burndown)")
    add_placeholder(doc, "Burndown Chart da Sprint 5")

    add_heading_custom(doc, "3. Sprint 6")
    add_heading_custom(doc, "3.1 Papéis Assumidos na Sprint 6")
    for item in [
        "Alyssandro Ramos: QA",
        "Rodrigo Almeida: Desenvolvedor e Documentação",
        "Djhonatah Wesley: Gerente de Projeto e Desenvolvedor",
        "José Lucas: Gerente de Projeto e Desenvolvedor",
    ]:
        add_bullet(doc, item)

    add_heading_custom(doc, "3.2 Backlog da Sprint")
    add_placeholder(
        doc,
        "Backlog da Sprint 6 (RF37–RF43 / quadro do gerenciador de projeto)",
    )

    add_heading_custom(doc, "3.3 Acompanhamento e Desempenho (Burndown)")
    add_placeholder(doc, "Burndown Chart da Sprint 6")

    out = RELEASES / "relatorio-processo-release3.docx"
    doc.save(out)
    return out


def gerar_release():
    doc = new_doc()
    capa(doc, "RELATÓRIO DE RELEASE - 3")

    add_heading_custom(doc, "1. Visão de Produto")
    add_body(
        doc,
        "O ClassRoomPB é uma aplicação de gestão acadêmica usada via terminal, "
        "direcionada à administração de cursos, disciplinas, períodos letivos, turmas e usuários "
        "distintos (administrador, coordenador, professor e aluno). A release 3 objetivou entregar "
        "o fluxo e as regras de negócio relacionadas a notas e avaliação, histórico acadêmico e "
        "relatórios exportáveis em PDF, complementando o ciclo acadêmico iniciado nas releases anteriores.",
    )

    add_heading_custom(doc, "2. Descrição Arquitetural")
    add_body(
        doc,
        "A aplicação adota arquitetura em camadas baseada em MVC estendido com "
        "repositórios. A camada de apresentação (view) concentra a interface CLI; os controllers "
        "orquestram regras de negócio e controle de acesso por perfil; o model encapsula entidades "
        "e invariantes; os repositories persistem o estado em arquivo JSON local. Nesta release "
        "foram acrescentados NotaController, HistoricoAcademicoController, RelatorioController e "
        "PdfRelatorioWriter. O fluxo operacional permanece: carregar dados → processar em memória → "
        "persistir após cada transação.",
    )

    add_heading_custom(doc, "2.1 Diagrama de Casos de Uso")
    add_body(
        doc,
        "Descrição visual da interação dos atores para com o sistema de acordo com suas "
        "capacidades, incluindo notas, histórico e relatórios da release 3.",
        first_line_indent=False,
    )
    add_image_if_exists(
        doc,
        DIAGRAMAS / "diagrama-casos-de-uso.png",
        width_cm=16,
        caption="Diagrama de Casos de Uso",
    )

    add_heading_custom(doc, "2.2 Diagrama de Classes")
    add_body(
        doc,
        "Visão descritiva das classes componentes do domínio do sistema, orquestração, "
        "interface e armazenamento, com destaque às entidades e controllers de notas, histórico e relatórios.",
        first_line_indent=False,
    )
    add_image_if_exists(
        doc,
        DIAGRAMAS / "diagrama-classes.png",
        width_cm=16,
        caption="Diagrama de Classes",
    )

    add_heading_custom(doc, "2.3 Diagramas de Sequência")
    for titulo, arquivos in SEQ_ITEMS:
        add_body(doc, titulo, first_line_indent=False)
        for nome in arquivos:
            add_image_if_exists(
                doc,
                SEQ / nome,
                width_cm=15.5,
                caption=nome.replace(".png", ""),
            )

    add_heading_custom(doc, "3. Funcionalidades Desenvolvidas na Release 3")
    for titulo, descricao, placeholder in FUNCIONALIDADES:
        add_bullet(doc, f"{titulo}: {descricao}")
        add_placeholder(doc, placeholder)

    add_heading_custom(doc, "4. Relatório de Testes e Cobertura de Código")
    add_table(
        doc,
        ["Indicador", "Valor"],
        [
            ("Data/hora", "20/07/2026 18:29:40"),
            ("Build", "SUCCESS"),
            ("Total de testes", "281"),
            ("Falhas", "0"),
            ("Erros", "0"),
            ("Ignorados", "0"),
            ("Cobertura total (instruções)", "85,0%"),
            ("Cobertura de branches", "66,0%"),
        ],
    )

    add_heading_custom(doc, "4.1 Resultado por Suite")
    suites = [
        ("MainTest", 2),
        ("controller.AutenticacaoControllerCadastroTest", 5),
        ("controller.AutenticacaoControllerLoginTest", 7),
        ("controller.ConfirmarMatriculaAutomaticaTest", 9),
        ("controller.CursoControllerTest", 12),
        ("controller.DisciplinaControllerTest", 12),
        ("controller.HistoricoRf36Rf39Test", 8),
        ("controller.ListaEsperaRf23Test", 10),
        ("controller.ListaEsperaRf24Test", 7),
        ("controller.ListaEsperaRf25Rf26Test", 6),
        ("controller.MatriculaControllerTest", 6),
        ("controller.MatriculaRequisitosRf16aRf22Test", 5),
        ("controller.NotaRf31Rf35Test", 15),
        ("controller.PeriodoLetivoControllerTest", 5),
        ("controller.PresencaRf27Test", 10),
        ("controller.PresencaRf28Test", 9),
        ("controller.PresencaRf29Rf30Test", 4),
        ("controller.RelatorioControllerTest", 10),
        ("controller.TurmaControllerTest", 12),
        ("model / repository / report / view (demais suites)", 127),
    ]
    add_table(
        doc,
        ["Suite de testes", "Testes", "Falhas", "Erros", "Ignorados"],
        [(nome, n, 0, 0, 0) for nome, n in suites] + [("Total", 281, 0, 0, 0)],
    )

    add_heading_custom(doc, "4.2 Cobertura por camada (JaCoCo)")
    add_body(
        doc,
        "Os controllers, models e repositories concentram a maior parte dos testes de "
        "unidade. A cobertura média total da aplicação ficou acima de 80%, atendendo o requisito "
        "estabelecido.",
        first_line_indent=False,
    )
    add_table(
        doc,
        ["Camada", "Cobertura", "Instr. cobertas", "Instr. não cobertas"],
        [
            ("Main (Classe primária)", "100,0%", "5", "0"),
            ("Controllers (regras de negócio)", "91,0%", "4.225", "384"),
            ("Models (entidades)", "88,0%", "1.961", "259"),
            ("Repositories (persistência)", "91,0%", "3.460", "305"),
            ("Report (PDF)", "91,0%", "126", "11"),
            ("View (CLI)", "73,0%", "3.777", "1.362"),
            ("Total da aplicação", "85,0%", "13.554", "2.321"),
        ],
    )

    out = RELEASES / "relatorio-release3.docx"
    doc.save(out)
    return out


def main():
    p1 = gerar_processo()
    p2 = gerar_release()
    print(p1)
    print(p2)


if __name__ == "__main__":
    main()
